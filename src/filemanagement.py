from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2txt
import os


def createDocument(quiz_data, question_bank = False):
    if os.path.exists("storage/output.docx"):
        os.remove("storage/output.docx")
    if os.path.exists("storage/output.pdf"):
        os.remove("storage/output.pdf")
    doc = Document()
    if question_bank:
        title = doc.add_heading('Question Bank', 1)
    else:
        title = doc.add_heading('Quiz', 1) 
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n")
    doc.add_heading("Questions : ", 2)
    for qno, q in enumerate(quiz_data):
        qno += 1
        para = doc.add_paragraph()
        para.add_run(f"{qno}) {q["question"]}\n\n").bold = True
        for choice, op in zip(q["options"], "abcd"):
            para.add_run(f"({op}) {choice}\n")
        if question_bank:
            para.add_run("\nCorrect Answer : ").bold = True
            para.add_run(q["answer"]+"\n")
            para.add_run("\nExplanation : \n").bold = True
            para.add_run(q["reason"]+"\n")
        para.add_run("\n")
    img = doc.add_picture('assets/watermark.png', width=Pt(150))
    img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save("storage/output.docx")

class ExtractContent:
    def __init__(self,file_path, file_type=None):
        self.__file_type = file_path.split(".")[-1]
        self.__file_path = file_path
        self.__words = 0
        self.__content = ""
        self.__pages = 0
        self.read()
    
    def getInformation(self):
        return f"FILE DESCRIPTION\nFile Path : {self.__file_path}\nNumber of Words : {self.__words}\nNumber of Pages : {self.__pages}"
    
    @staticmethod
    def readPdf(file_path):
        reader = PdfReader(file_path)
        pages = len(reader.pages)
        content = ""
        for i in range(pages):
            page = reader.pages[i]
            text = page.extract_text()
            content += text
        words = len(content.split())
        content = content.replace("\n"," ")
        return content, words, pages
    
    @staticmethod
    def readText(file_path):
        content = ""
        with open(file_path,"r") as f:
            content = f.read()
            f.close()
        words = len(content.split())
        return content, words, 1
        
    @staticmethod
    def readDocx(file_path):
        doc = Document(file_path)
        content = ""
        for para in doc.paragraphs:
            content+=para.text
        words = len(content.split())
        return content, words, 1
    
    def read(self):
        if self.__file_type == "pdf":
            self.__content,self.__words, self.__pages  = ExtractContent.readPdf(self.__file_path)
        elif self.__file_type == "txt":
            self.__content,self.__words, self.__pages  = ExtractContent.readText(self.__file_path)
        elif self.__file_type == "docx":
            self.__content,self.__words, self.__pages  = ExtractContent.readDocx(self.__file_path)
        else:
            raise Exception("Only .pdf, .txt, .docx files are allowed")
    
    def getContent(self):
        return self.__content